# -*- coding: UTF-8 -*-
import sys
import math
import re

from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from xml.dom.minidom import parseString
import xml.etree.ElementTree as ET
from lxml import etree

from docx.opc.constants import RELATIONSHIP_TYPE as RT
reload(sys)
sys.setdefaultencoding('utf8')

def get_analyze_the_document(path):
    file_analyzes = {}
    def get_general_properties(document):
        general_properties = {}
        general_properties["author"] = str(document.core_properties.author) # Автор документа
        general_properties["created"] = str(document.core_properties.created) # Дата создания документа
        general_properties["last_modified_by"] = str(document.core_properties.last_modified_by) # Пользователь, который менял последний
        general_properties["modified"] = str(document.core_properties.modified) # Время изменения
        general_properties["title"] = str(document.core_properties.title) #Название
        file_analyzes["general_properties"] = general_properties

    def get_document_margins(document):
        margins = {}
        sections = document.sections
        margins["top"] = str(round(sections[0].top_margin.cm, 2))
        margins["bottom"] = str(round(sections[0].bottom_margin.cm, 2))
        margins["left"] = str(round(sections[0].left_margin.cm, 2))
        margins["right"] = str(round(sections[0].right_margin.cm, 2))
        file_analyzes["margins"] = margins

    def get_headers_texts(document):
        document_headers_texts = []
        count = 0
        for p in document.paragraphs:
            if p.style.name != "Normal":
                if "Заголовок" in p.style.name and "ФИО" in p.style.name:
                    text = re.sub(r'\s+', ' ', p.text)
                    document_headers_texts.append((text))
            if 'toc' in p.style.name:
                count +=1

        file_analyzes["menu_item_count"] = count
        file_analyzes["document_headers_texts"] = document_headers_texts

    def get_subject_index(document):
        subject_index_temp = []
        subject_index = []
        fi = 0
        all_p = document.paragraphs
        for i, val in enumerate(all_p):
            if "Заголовок" in val.style.name and "предметный указатель" in val.text.lower():
                fi = i
        all_p = all_p[(fi+1):]

        for p in all_p:
            if "указатель" in p.style.name.lower():
                subject_index_temp.append(p.text)
        for s in subject_index_temp:
            i = s.rfind(',')
            subject_index.append(re.sub(r'\s+', ' ', s[:i]))
        file_analyzes["subject_index"] = subject_index

    def get_custom_styles(document):
        custom_styles = {}
        all_docx_styles = document.styles
        for s in all_docx_styles:
            if s.builtin == False and "ФИО" in s.name:
                if "Заголовок" in s.name:
                    header_style = {}
                    custom_head_style = all_docx_styles[s.base_style.name]
                    header_style["name"] = str(s.name)
                    header_style["font_name"] = str(s.font.name)
                    header_style["font_size"] = str(s.font.size)
                    header_style["font_italic"] = str(s.font.italic)
                    header_style["font_bold"] = str(s.font.bold)
                    header_style["line_spacing"] = str(s.paragraph_format.line_spacing)
                    header_style["first_line_indent"] = str(s.paragraph_format.first_line_indent)
                    header_style["space_before"] = str(s.paragraph_format.space_before)
                    header_style["space_after"] = str(s.paragraph_format.space_after)
                    header_style["alignment"] = str(s.paragraph_format.alignment)
                    custom_styles["header_style"] = str(header_style)
                else:
                    paragraph_style = {}
                    custom_paragraph_style = all_docx_styles[s.base_style.name]
                    header_style["name"] = s.name
                    header_style["font_name"] = s.font.name

                    # print s.paragraph_format.alignment

        file_analyzes["custom_styles"] = custom_styles

    def get_document_numbering(document):
        def iter_header_parts(document):
            document_part = document.part
            for rel in document_part.rels.values():
                if rel.reltype == RT.FOOTER:
                    yield rel.target_part

        document_page_numbering = False
        for header_part in iter_header_parts(document):

            header_xml = header_part._blob
            namespace = dict(w="http://schemas.openxmlformats.org/wordprocessingml/2006/main")
            root = ET.fromstring(header_xml)
            # print header_xml
            try:
                text_element = root.find(".//w:docPartGallery", namespace)
                if "Page Numbers" in ET.tostring(text_element):
                    document_page_numbering = True
            except:
                pass
        file_analyzes["document_page_numbering"] = document_page_numbering

    def get_document_heading(document):
        document_header = ""
        def iter_header_parts(document):
            document_part = document.part
            for rel in document_part.rels.values():
                if rel.reltype == RT.HEADER:
                    yield rel.target_part


        for header_part in iter_header_parts(document):
            header_xml = header_part._blob
            namespace = dict(w="http://schemas.openxmlformats.org/wordprocessingml/2006/main")
            root = ET.fromstring(header_xml)
            text_element = root.find(".//w:t", namespace)
            if text_element is not None:
                if text_element.text != '':
                    document_header = text_element.text
        file_analyzes["document_header"] = document_header

    document = Document(path)
    get_headers_texts(document)
    get_general_properties(document)
    get_document_margins(document)
    get_custom_styles(document)
    get_subject_index(document)
    get_document_numbering(document)
    get_document_heading(document)
    return file_analyzes

correct_analyze_object = get_analyze_the_document('data/internet.docx')
student_analyze_object = get_analyze_the_document('data/internet dummy.docx')


def documents_comments(correct_analyze_object, student_analyze_object):
    list = []
    cao = correct_analyze_object
    sao = student_analyze_object

    if (cao["margins"]["top"] == sao["margins"]["top"]):
        list.append("Верхний отступ заполен верно")
    else:
        list.append("Ошибка при заполнии верхнего отступа")

    if (cao["margins"]["bottom"] == sao["margins"]["bottom"]):
        list.append("Нижний отступ заполен верно")
    else:
        list.append("Ошибка при заполнии нижнего отступа")

    if (cao["margins"]["left"] == sao["margins"]["left"]):
        list.append("Левый отступ заполен верно")
    else:
        list.append("Ошибка при заполнии левого отступа")

    if (cao["margins"]["right"] == sao["margins"]["right"]):
        list.append("Правый отступ заполен верно")
    else:
        list.append("Ошибка при заполнии правого отступа")

    about_headers_msg = ""
    headers_diff =  [item for item in cao["document_headers_texts"] if item not in sao["document_headers_texts"]]
    if(len(headers_diff) == 0):
        labout_headers_msg = "Все заголовки оформлены верно"
    else:
        about_headers_msg = "Некоторые заголовки оформлены неверно, обратите внимание на заголовки: "
        for diff in headers_diff:
            about_headers_msg += diff + " "
    list.append(about_headers_msg)


    # ДОДЕЛАТЬ МЕНЮ
    menu_msg = ""
    if(cao["menu_item_count"] == sao["menu_item_count"]):
        menu_msg = "Меню создано"
    else:
        menu_msg = "Меню не создано"
    list.append(menu_msg)

    about_subject_index_msg = ""
    subject_index_diff =  [item for item in cao["subject_index"] if item not in sao["subject_index"]]
    if(len(subject_index_diff) == 0):
        about_subject_index_msg = "Предметный указатель создан верно"
    else:
        about_subject_index_msg = "Предметный указатель создан неверно, обратите внимание на: "
        for diff in subject_index_diff:
            about_subject_index_msg += diff + " "
    list.append(about_subject_index_msg)

    page_num_msg = ""
    if(cao["document_page_numbering"] == sao["document_page_numbering"]):
        page_num_msg = "Нумерация страниц выставлена верно"
    else:
        page_num_msg = "Нумерация страниц выставлена не верно"
    list.append(page_num_msg)

    doc_heading_msg = ""
    if(cao["document_header"] == sao["document_header"]):
        doc_heading_msg = "Верхний колонтитул заполнен верно"
    else:
        doc_heading_msg = "Верхний колонтитул заполнен не верно"
    list.append(doc_heading_msg)

    return list


coments = documents_comments(correct_analyze_object, student_analyze_object)

print correct_analyze_object

for coment in coments:
    print coment

# document = Document('data/internet dummy.docx')
#
# for p in document.paragraphs:
#     print p.style.name, p.text
        # if text_element is not None:
    #     if text_element.text != '':
    #         pass
    #         # print(text_element.text)

#
# tables = document.tables
#
# for sec in document.sections:
#      print sec.header.is_linked_to_previous

# for row in tables[0].rows:
#     for cell in row.cells:
#         for paragraph in cell.paragraphs:
            # print paragraph.text, paragraph.alignment


# for row in tables[0].rows:
#     print row.cells[0].paragraphs[0].text + row.cells[2].paragraphs[0].text

# get_section_margins(document.sections)
# print  file_analyzes

# for s in document.paragraphs:
#      print(s.style.name)

# body_element = document._body._body
# print(body_element.xml)
#
# ps = body_element.xpath('./w:hyperlink/w:r')
#
# paragraphs = [Paragraph(p, None) for p in ps]
#



# print "!!!", ps

# styles = document.styles
# for style in styles:
#     if "ФИО" in style.name:
#         try:
#             print style.font.name
#         except:
#             print 'Родительский стиль: ', style.base_style.name
#         print '\n'

#style.builtin #Стиль сделан пользователем
#style.paragraph_format.first_line_indent.cm # Отступ первой строки
#style.paragraph_format.space_before.pt # Отступ до
#style.paragraph_format.space_after.pt # Отступ после
#style.paragraph_format.line_spacing # Межстрочный интервал